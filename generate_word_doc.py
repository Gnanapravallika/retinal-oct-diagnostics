# =========================================================================
# Retinal OCT Literature Review: Word Document Generator
# =========================================================================
# This script automatically installs 'python-docx' if missing and generates
# a beautifully formatted MS Word document (retinal_oct_literature_review.docx)
# containing all 28 structured research tables for your thesis.
# =========================================================================

import os
import sys

# Step 1: Programmatically verify and install python-docx if missing
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Installing 'python-docx' library for Word document generation...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background shading of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxas (twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_table(doc, headers, rows):
    """Creates a beautifully padded, shaded academic-style table."""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.autofit = False
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = headers[0]
    hdr_cells[1].text = headers[1]
    
    for cell in hdr_cells:
        set_cell_background(cell, "2C3E50") # Deep charcoal
        set_cell_margins(cell, top=120, bottom=120)
        p = cell.paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            
    # Format Data Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        row_cells[0].text = str(row_data[0])
        row_cells[1].text = str(row_data[1])
        
        # Alternating background shading
        bg_color = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
        
        for c_idx, cell in enumerate(row_cells):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                # Bold the question column
                if c_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(44, 62, 80)
                else:
                    run.font.color.rgb = RGBColor(60, 60, 60)
                    
    # Adjust widths (Column 1 narrower than Column 2)
    for row in table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)
        
    doc.add_paragraph() # Spacing

def main():
    print("Generating 'retinal_oct_literature_review.docx'...")
    doc = docx.Document()
    
    # Document Title
    title = doc.add_paragraph()
    run = title.add_run("Retinal OCT Classification & Explainability\nComprehensive Literature Review")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    run.font.name = 'Arial'
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # SECTION 0: Research Mentor's Literature Closure & Guidelines
    doc.add_heading("Part 1: Research Mentoring & Literature Review Closure", level=1)
    
    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "Based on the rigorous technical and medical AI requirements of 2026, the literature review phase is officially frozen. "
        "Adding more papers yields diminishing returns; instead, efforts are directed toward presenting the scientific findings and "
        "experimental validations. Below is the finalized literature scope, structure, research gaps, and contributions defining the paper's story.\n\n"
    ).font.name = 'Arial'
    
    # Final Literature Scope Table
    doc.add_heading("Final Literature Scope", level=2)
    add_styled_table(doc, ["Literature Section Theme", "Recommended Key References count"], [
        ("Retinal OCT Classification", "6-8 key references"),
        ("Attention Mechanisms", "5-6 key references"),
        ("Explainable AI (XAI)", "5-6 key references"),
        ("Faithfulness Evaluation", "5-6 key references"),
        ("External Validation & Generalization", "5-6 key references"),
        ("Clinical AI & Reporting Calibration", "4-5 key references"),
        ("Total References in Paper", "≈45-60 references (30-35 key review papers)")
    ])
    
    # Final Literature Structure
    doc.add_heading("Final Literature Structure (Manuscript Chapter 2)", level=2)
    
    p_struct = doc.add_paragraph()
    r_struct = p_struct.add_run(
        "2.1 Deep Learning for Retinal OCT Classification\n"
        "   - References: Kermany et al., De Fauw et al., ResNet, OCTDL, DenseNet, EfficientNet.\n"
        "   - Purpose: Establish the clinical problem, explain why CNNs became dominant, and highlight remaining challenges.\n\n"
        "2.2 Attention Mechanisms for OCT Analysis\n"
        "   - References: CBAM, GABNet, Lesion-Aware CNN, Multi-Level Dual Attention, Vision Transformer.\n"
        "   - Purpose: Explain channel/spatial attention and discuss strengths and limitations.\n\n"
        "2.3 Explainable AI for Medical Imaging\n"
        "   - References: Grad-CAM, LayerCAM, Score-CAM, Integrated Gradients, ReLaX.\n"
        "   - Purpose: Compare explanation methods and explain why qualitative visualization alone is insufficient.\n\n"
        "2.4 Quantitative Evaluation of Explainability\n"
        "   - References: Samek, ROAD, ROAR, Quantus, Sanity Checks.\n"
        "   - Purpose: Motivate AOPC, deletion/insertion tests, and justify objective faithfulness assessments.\n\n"
        "2.5 External Validation and Domain Generalization\n"
        "   - References: External Validation survey, DomainBed, RETFound, MedViT external validation, Scanner intensity shift.\n"
        "   - Purpose: Explain why internal accuracy is not enough and motivate independent test domain validation.\n\n"
        "2.6 Clinical AI Reliability\n"
        "   - References: TRIPOD-AI, CONSORT-AI, DECIDE-AI, Guo et al. (Calibration).\n"
        "   - Purpose: Discuss transparency, reproducibility, ECE calibration, and responsible reporting."
    )
    r_struct.font.name = 'Arial'
    r_struct.font.size = Pt(10)
    r_struct.font.color.rgb = RGBColor(60, 60, 60)
    
    # Final Research Gap
    doc.add_heading("Final Research Gap Statement", level=2)
    p_gap = doc.add_paragraph()
    r_gap = p_gap.add_run(
        "\"Deep learning has substantially improved retinal OCT classification, and attention mechanisms have enhanced feature representation. "
        "However, many studies emphasize predictive performance while relying primarily on qualitative explanation methods and internal test sets. "
        "Comparatively fewer studies combine quantitative faithfulness evaluation with external validation to assess whether attention-guided models "
        "remain reliable across independent datasets. Motivated by this gap, this study investigates an attention-enhanced ResNet framework using "
        "quantitative explainability analysis and external validation within a reproducible evaluation pipeline.\""
    )
    r_gap.font.name = 'Arial'
    r_gap.font.italic = True
    r_gap.font.bold = True
    r_gap.font.color.rgb = RGBColor(44, 62, 80)
    
    # Final Contributions
    doc.add_heading("Final Contributions", level=2)
    p_cont = doc.add_paragraph()
    r_cont = p_cont.add_run(
        "1. An attention-enhanced ResNet-50 framework for multi-class retinal OCT classification using patient-level data partitioning.\n"
        "2. A quantitative explainability evaluation using LayerCAM with perturbation-based faithfulness metrics (e.g., AOPC and related deletion/insertion analyses).\n"
        "3. External validation on an independent OCT dataset to assess generalization.\n"
        "4. A reproducible experimental pipeline with structured preprocessing, baseline comparisons, ablation studies, and transparent reporting."
    )
    r_cont.font.name = 'Arial'
    r_cont.font.size = Pt(10)
    r_cont.font.color.rgb = RGBColor(44, 62, 80)

    # Final Story
    doc.add_heading("Final Story of the Paper", level=2)
    p_story = doc.add_paragraph()
    r_story = p_story.add_run(
        "Retinal OCT Diagnosis\n"
        "        \u2514\u2500\u2500\u25BA Attention-Enhanced Feature Learning\n"
        "                \u2514\u2500\u2500\u25BA Explainability\n"
        "                        \u2514\u2500\u2500\u25BA Quantitative Faithfulness\n"
        "                                \u2514\u2500\u2500\u25BA External Validation\n"
        "                                        \u2514\u2500\u2500\u25BA Trustworthy Clinical AI\n\n"
        "Note: Attention is not the endpoint. It is one component of a broader story about trustworthy evaluation."
    )
    r_story.font.name = 'Consolas'
    r_story.font.size = Pt(10)
    r_story.font.color.rgb = RGBColor(80, 80, 80)

    # Literature status
    doc.add_heading("Literature Status Checklist", level=2)
    add_styled_table(doc, ["Task / Milestone", "Status"], [
        ("Core papers review", "COMPLETED"),
        ("Advanced papers review", "COMPLETED"),
        ("Literature themes mapping", "FINALIZED"),
        ("Research gap formulation", "FINALIZED (working draft)"),
        ("Paper positioning", "FINALIZED"),
        ("Literature Scope", "FROZEN (Officially Closed)")
    ])
    
    doc.add_paragraph()
    
    # SECTION 1.5: Phase 2 - Research Design
    doc.add_heading("Part 2: Research Design & Experimental Specifications", level=1)
    
    p_rd = doc.add_paragraph()
    r_rd = p_rd.add_run(
        "To ensure complete academic rigor and prevent random ad-hoc coding, the research design is frozen. "
        "Every code execution and table output must align directly with the following specifications:\n\n"
        "1. RESEARCH QUESTION (FINAL):\n"
        "Can attention-guided feature learning improve both retinal OCT classification performance and explanation faithfulness "
        "while maintaining robust generalization under external validation?\n\n"
        "2. RESEARCH OBJECTIVES:\n"
        "  - Objective 1: Develop an attention-enhanced ResNet-50 framework for multi-class retinal OCT classification.\n"
        "  - Objective 2: Quantitatively evaluate explanation faithfulness.\n"
        "  - Objective 3: Assess external generalization.\n"
        "  - Objective 4: Analyze model reliability."
    )
    r_rd.font.name = 'Arial'
    
    # Hypotheses
    doc.add_heading("Testable Hypotheses", level=2)
    add_styled_table(doc, ["Hypothesis ID", "Theoretical & Empirical Claim"], [
        ("H1", "AE-ResNet significantly improves Macro F1 over the baseline ResNet-50."),
        ("H2", "Attention-guided learning produces more faithful explanations than the baseline (measured via AOPC Deletion/Insertion)."),
        ("H3", "External validation causes a performance decrease, but AE-ResNet maintains better performance than baseline."),
        ("H4", "Attention-guided features remain spatially consistent under domain shift.")
    ])
    
    # Variables
    doc.add_heading("Experimental Variables", level=2)
    add_styled_table(doc, ["Variable Category", "Isolated Experimental Parameters"], [
        ("Independent Variable", "Model Architecture (ResNet-50, DenseNet-121, EfficientNet-B0, AE-ResNet)"),
        ("Dependent Variables (Classification)", "Accuracy, Precision, Recall, Macro F1, ROC-AUC"),
        ("Dependent Variables (Explainability)", "LayerCAM, Deletion test, Insertion test, AOPC"),
        ("Dependent Variables (Reliability)", "External validation accuracy shift, Calibration error (ECE)")
    ])
    
    # Pipeline
    doc.add_heading("Experimental Pipeline Sequence", level=2)
    p_pipe = doc.add_paragraph()
    r_pipe = p_pipe.add_run(
        "OCTDL Dataset \u2192 Patient-Level Split \u2192 Preprocessing \u2192 Augmentation \u2192 Baseline Models \u2192 AE-ResNet \u2192 Training \u2192 Testing \u2192 LayerCAM \u2192 AOPC \u2192 Deletion \u2192 Insertion \u2192 External Validation \u2192 Statistics \u2192 Paper Tables"
    )
    r_pipe.font.name = 'Arial'
    r_pipe.font.bold = True
    r_pipe.font.color.rgb = RGBColor(44, 62, 80)
    
    # Baselines & Ablation Tables
    doc.add_heading("Baselines & Ablation Setup", level=2)
    add_styled_table(doc, ["Model Baseline", "Purpose / Context in Medical Imaging"], [
        ("ResNet-50", "Control standard residual CNN baseline"),
        ("DenseNet-121", "Ophthalmic imaging comparison standard"),
        ("EfficientNet-B0", "Computationally efficient CNN comparison baseline"),
        ("AE-ResNet", "Proposed attention-enhanced architecture")
    ])
    
    add_styled_table(doc, ["Ablation Step", "Component Isolated"], [
        ("1. Baseline", "Standard ResNet-50"),
        ("2. + Weighted Loss", "Effect of loss cost balancing"),
        ("3. + Weighted Sampler", "Effect of oversampling minority classes"),
        ("4. + Channel Attention", "Effect of Channel Attention ('What' stream)"),
        ("5. + Spatial Attention", "Effect of Spatial Attention ('Where' stream)"),
        ("6. Full AE-ResNet", "Unified proposed architecture")
    ])
    
    # Statistical Analysis
    doc.add_heading("Statistical & Evaluation Protocol", level=2)
    p_stat = doc.add_paragraph()
    r_stat = p_stat.add_run(
        "For every dependent variable, we evaluate over n = 5 independent seeds. "
        "Results are reported as Mean \u00b1 Standard Deviation (SD) alongside 95% Confidence Intervals (CI). "
        "Wilcoxon signed-rank tests are computed to determine the statistical significance (p < 0.01) of "
        "attributions compared to random baselines."
    )
    r_stat.font.name = 'Arial'
    
    # Expected Tables & Figures
    doc.add_heading("Expected Manuscript Artifacts", level=2)
    add_styled_table(doc, ["Expected Tables", "Expected Figures"], [
        ("Table 1: Dataset statistics (prevalence & split ratios)", "Figure 1: Complete experimental pipeline block diagram"),
        ("Table 2: Baseline comparisons (Acc, F1, ECE, AOPC)", "Figure 2: Detailed AE-ResNet architecture diagram"),
        ("Table 3: Ablation study details", "Figure 3: Training & validation loss curves"),
        ("Table 4: Internal evaluation (OCTDL test metrics)", "Figure 4: Confusion matrix on target classes"),
        ("Table 5: External evaluation (OCTID un-normalized vs normalized)", "Figure 5: ROC-AUC curves for multi-class diagnostics"),
        ("Table 6: Explainability metrics (AOPC Deletion vs Insertion)", "Figure 6: LayerCAM visual overlays on retinal tissues"),
        ("Table 7: Statistical summary (significance & CIs)", "Figure 7 & 8: Deletion & Insertion perturbation curves"),
        ("-", "Figure 9: External validation saliency drift visualization")
    ])

    # Balanced Notebook Structure
    doc.add_heading("M.Tech Thesis Execution Plan (Balanced Notebook Structure)", level=2)
    p_nb = doc.add_paragraph()
    r_nb = p_nb.add_run(
        "To balance MLOps rigor with execution speed and reduce overhead, we utilize a single main notebook "
        "(AE_ResNet_Project.ipynb) containing structured execution cells, backed by a modular 'src/' Python package. "
        "This keeps implementation files separate from experiment logs, avoiding giant unreadable code cells.\n\n"
        "The notebook contains 19 sequential research cells:\n"
        "  - Cell 1: Setup & libraries initialization\n"
        "  - Cell 2: Dataset loading & validation\n"
        "  - Cell 3: Dataset stats analysis (Outputs: Table 1)\n"
        "  - Cell 4: Patient-level split validation (Outputs: Patient leakage audit)\n"
        "  - Cell 5: DataLoader & sampler initialization\n"
        "  - Cell 6: Preprocessing visual checks (Outputs: Figure 1)\n"
        "  - Cell 7: Baseline models setup (ResNet, DenseNet, EfficientNet)\n"
        "  - Cell 8: AE-ResNet architecture loading\n"
        "  - Cell 9: Model training loop executions\n"
        "  - Cell 10: Model evaluation metrics compilation (Outputs: Table 2)\n"
        "  - Cell 11: Ablation test executions (Outputs: Table 3)\n"
        "  - Cell 12: LayerCAM activations hook extraction\n"
        "  - Cell 13: Deletion test execution (Outputs: Figure 7)\n"
        "  - Cell 14: Insertion test execution (Outputs: Figure 8)\n"
        "  - Cell 15: AOPC score calculations (Outputs: Table 6)\n"
        "  - Cell 16: External Validation audit on OCTID (Outputs: Table 5, Figure 9)\n"
        "  - Cell 17: Multi-seed statistical validations (Outputs: Table 7)\n"
        "  - Cell 18: Publication-ready figure plotting\n"
        "  - Cell 19: Save all assets to results/ folder"
    )
    r_nb.font.name = 'Arial'
    r_nb.font.size = Pt(10)
    r_nb.font.color.rgb = RGBColor(60, 60, 60)
    
    doc.add_paragraph()

    # SECTION 1.6: Phase 3 - Research Implementation Blueprint
    doc.add_heading("Part 3: Research Implementation Blueprint & Repository Structure", level=1)
    
    p_rib = doc.add_paragraph()
    r_rib = p_rib.add_run(
        "To shift the study from describing a model to providing reproducible experimental evidence, "
        "every execution component is structured around a scientific question. We use a single main Colab notebook "
        "for the experimental workflow, combined with a modular 'src/' package for implementation. "
        "This minimizes switching overhead while maintaining software engineering rigor.\n\n"
        "1. THE EVIDENCE-FIRST BLUEPRINT:\n"
        "  - Cell 2-5: Is the dataset suitable and free from patient leakage? (Outputs: Table 1)\n"
        "  - Cell 6: Does preprocessing enhance tissue contrast without distorting anatomy? (Outputs: Figure 1)\n"
        "  - Cell 7: How strong are standard, control deep classifiers? (Outputs: Table 2)\n"
        "  - Cell 8-10: Does attention-gated representation improve performance? (Outputs: Table 2)\n"
        "  - Cell 11: Which components of the network drive the improvement? (Outputs: Table 3)\n"
        "  - Cell 12-15: Are the visual explanations mathematically faithful? (Outputs: Table 6, Figure 7, 8)\n"
        "  - Cell 16: Does the model generalize to unseen scanner hardware? (Outputs: Table 5, Figure 9)\n"
        "  - Cell 17: Are the observed improvements statistically consistent? (Outputs: Table 7)\n"
        "  - Cell 18-19: Compilation of all publication-ready figures and tables."
    )
    r_rib.font.name = 'Arial'
    
    # Final Repository Structure
    doc.add_heading("Final Academic Repository Structure", level=2)
    p_repo = doc.add_paragraph()
    r_repo = p_repo.add_run(
        "ae_resnet_oct/\n"
        "\u251c\u2500\u2500 AE_ResNet_Project.ipynb   \u2190 ONE main notebook\n"
        "\u251c\u2500\u2500 src/\n"
        "\u2502   \u251c\u2500\u2500 dataset/\n"
        "\u2502   \u251c\u2500\u2500 preprocessing/\n"
        "\u2502   \u251c\u2500\u2500 models/\n"
        "\u2502   \u251c\u2500\u2500 training/\n"
        "\u2502   \u251c\u2500\u2500 evaluation/\n"
        "\u2502   \u2514\u2500\u2500 xai/\n"
        "\u251c\u2500\u2500 results/\n"
        "\u2502   \u251c\u2500\u2500 checkpoints/\n"
        "\u2502   \u251c\u2500\u2500 metrics/\n"
        "\u2502   \u251c\u2500\u2500 figures/\n"
        "\u2502   \u2514\u2500\u2500 tables/\n"
        "\u251c\u2500\u2500 paper/\n"
        "\u251c\u2500\u2500 tests/\n"
        "\u2514\u2500\u2500 requirements.txt\n"
    )
    r_repo.font.name = 'Consolas'
    r_repo.font.size = Pt(10)
    r_repo.font.color.rgb = RGBColor(60, 60, 60)
    doc.add_page_break()

    # SECTION A: Foundational Papers
    doc.add_heading("Section 3: Foundational Retinal AI Papers", level=1)



    
    # 1. Kermany et al. (2018)
    doc.add_heading("1. Kermany et al. (2018)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Automated classification of treatable blinding retinal diseases to improve referral workflows."),
        ("Dataset(s)", "UCSD cohort (108,309 OCT scans from 4,686 patients)."),
        ("Model", "InceptionV3."),
        ("Explainability used?", "Yes, coarse bounding boxes generated via spatial occlusion testing."),
        ("External validation?", "Yes, evaluated on 1,000 independent images from a separate clinic."),
        ("Strengths", "Massive patient cohort; direct benchmarking against human experts; published in Cell."),
        ("Weaknesses", "Occlusion maps are highly coarse, blocky, and fail to localize fine layer structures."),
        ("What can our paper learn?", "Highlights the need for pixel-level visual attributions (LayerCAM) over coarse occlusion boxes.")
    ])

    # 2. De Fauw et al. (2018)
    doc.add_heading("2. De Fauw et al. (2018)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Automated triage and referral recommendation for over 50 sight-threatening retinal pathologies."),
        ("Dataset(s)", "Moorfields Eye Hospital London cohort (14,884 3D scans)."),
        ("Model", "Two-stage framework: 3D U-Net (segmentation) + 3D CNN (classification/referral)."),
        ("Explainability used?", "Yes, intermediate tissue segmentation maps act as the clinical explanation."),
        ("External validation?", "Yes, validated on scans from a different scanner manufacturer."),
        ("Strengths", "Volumetric inputs; matches clinic workflow; published in Nature Medicine."),
        ("Weaknesses", "Requires massive manual pixel-level segmentation labeling to train the U-Net."),
        ("What can our paper learn?", "Clinicians trust layer localization. We use spatial attention to constrain representations without segmentations.")
    ])

    # 3. He et al. (2016) - ResNet
    doc.add_heading("3. ResNet (He et al., 2016)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Vanishing gradients during optimization in deep neural networks."),
        ("Dataset(s)", "ImageNet, CIFAR-10, PASCAL VOC."),
        ("Model", "ResNet-34/50/101/152 utilizing identity shortcut connections."),
        ("Explainability used?", "No (pure black box)."),
        ("External validation?", "Yes, validated across downstream transfer tasks."),
        ("Strengths", "Solved deep gradient propagation; standard industry baseline backbone."),
        ("Weaknesses", "Treats background noise and pathological features with equal weight due to a lack of feature-selectivity."),
        ("What can our paper learn?", "Use ResNet-50 as our residual feature extractor, but inject attention modules (CSA) to filter noise.")
    ])

    # 4. CBAM (Woo et al., 2018)
    doc.add_heading("4. CBAM (Woo et al., 2018)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Standard CNN layers treat all channels and spatial pixels equally, ignoring relative feature importance."),
        ("Dataset(s)", "ImageNet, MS COCO, VOC."),
        ("Model", "ResNet/ResNeXt baselines integrated with sequential Channel and Spatial Attention gates."),
        ("Explainability used?", "Yes, qualitative Grad-CAM heatmaps."),
        ("External validation?", "Yes (on benchmark test splits)."),
        ("Strengths", "Lightweight, modular attention; sequential channel-first design out-performed parallel designs."),
        ("Weaknesses", "Saliency evaluation was purely subjective. Did not test explanation faithfulness under domain shifts."),
        ("What can our paper learn?", "Sequential CBAM/CSA attention is highly effective. We must quantitatively verify its faithfulness using AOPC.")
    ])

    # 5. GABNet (Huang et al., 2023)
    doc.add_heading("5. GABNet (Huang et al., 2023)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Efficient and lightweight retinal disease classification from OCT scans."),
        ("Dataset(s)", "UCSD dataset."),
        ("Model", "GABNet (integrates a Global Attention Block weighting height, width, and channel)."),
        ("Explainability used?", "Yes (qualitative Grad-CAM)."),
        ("External validation?", "No (evaluated only on in-distribution UCSD test split)."),
        ("Strengths", "Efficient, lightweight attention module suitable for constrained hardware."),
        ("Weaknesses", "Lacks domain shift validation; evaluation is limited to a single clean dataset; explainability is purely visual."),
        ("What can our paper learn?", "Proves attention is useful for OCT, but reveals a gap: they did not test model generalization or explainability under shifts.")
    ])

    # 6. Grad-CAM (Selvaraju et al., 2017)
    doc.add_heading("6. Grad-CAM (Selvaraju et al., 2017)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Lack of interpretability in deep learning model predictions."),
        ("Dataset(s)", "ImageNet, PASCAL VOC."),
        ("Model", "VGG, ResNet, AlexNet."),
        ("Explainability used?", "Yes, proposed Grad-CAM (gradients of target scores with respect to terminal conv layers)."),
        ("External validation?", "Evaluated via human trust surveys."),
        ("Strengths", "Model-agnostic; requires no retraining."),
        ("Weaknesses", "Gradients are averaged globally, producing coarse, blurry heatmaps that cannot locate fine retinal layer structures."),
        ("What can our paper learn?", "Grad-CAM is a good baseline, but we must use LayerCAM to capture fine-grained pixel attributions, and verify them.")
    ])

    # 7. Samek et al. (2017)
    doc.add_heading("7. Samek et al. (2017)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Saliency maps can look visually appealing to humans but fail to represent what the model actually computed."),
        ("Dataset(s)", "MNIST, SUN, ILSVRC."),
        ("Model", "Feed-forward CNNs."),
        ("Explainability used?", "Yes, proposed the Deletion/Perturbation Test to calculate AOPC."),
        ("External validation?", "N/A (methodological paper)."),
        ("Strengths", "Replaced subjective human evaluations with a rigorous mathematical benchmark."),
        ("Weaknesses", "Did not investigate how explainability maps behave or degrade when the model encounters domain distribution shifts."),
        ("What can our paper learn?", "Adopt the AOPC deletion and insertion formulas as our primary quantitative explainability evaluation metrics.")
    ])

    # 8. OCTDL Dataset Paper (Jiang et al., 2024)
    doc.add_heading("8. OCTDL Dataset Paper (Jiang et al., 2024)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Scarcity of public, high-resolution, multi-class OCT datasets containing rare pathologies."),
        ("Dataset(s)", "OCTDL (2,821 scans across 7 diagnostic classes)."),
        ("Model", "Standard ResNet, DenseNet, and EfficientNet baselines."),
        ("Explainability used?", "No."),
        ("External validation?", "No."),
        ("Strengths", "Covers 7 classes including rare vascular pathologies (RAO, VID)."),
        ("Weaknesses", "Severe class imbalance; baseline models do not evaluate domain generalization, calibration, or explainability."),
        ("What can our paper learn?", "Use this dataset as our training domain, address its severe imbalance, and solve the generalization and explainability gaps.")
    ])

    # SECTION B: Advanced Skimmed Papers
    doc.add_page_break()
    doc.add_heading("Section 2: Skimmed Retinal AI Papers", level=1)

    # 9. Attention to Lesion (2020)
    doc.add_heading("9. Attention to Lesion (LACNN, 2020)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Standard CNNs extract global features, missing local pathological lesions."),
        ("Dataset(s)", "Private clinical dataset."),
        ("Model", "Lesion-Aware CNN (LACNN) integrating a lesion-detection branch."),
        ("Explainability used?", "Yes (lesion localization maps)."),
        ("External validation?", "No."),
        ("Strengths", "Focuses representations directly on pathological biomarkers (lesions)."),
        ("Weaknesses", "Requires bounding box annotations of lesions for training, which has high labeling costs."),
        ("What can our paper learn?", "Focus on pathology biomarkers is vital. Since we don't have bounding boxes, we use spatial attention (CSA) to lock layers.")
    ])

    # 10. Multi-Level Dual-Attention CNN (2021)
    doc.add_heading("10. Multi-Level Dual-Attention CNN (2021)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Multi-scale macular disease feature extraction is hindered by varied lesion sizes."),
        ("Dataset(s)", "UCSD dataset."),
        ("Model", "Multi-Level Dual-Attention CNN."),
        ("Explainability used?", "Yes (qualitative Grad-CAM)."),
        ("External validation?", "No."),
        ("Strengths", "Integrates spatial and channel attention across multiple network levels."),
        ("Weaknesses", "Lacks cross-scanner validation and quantitative verification of the attention maps."),
        ("What can our paper learn?", "Sequential multi-level attention gating is powerful, supporting our choice of injecting CSA gates in Layers 3 and 4.")
    ])

    # 11. Interpretable Transformer Network (2021)
    doc.add_heading("11. Interpretable Transformer Network (2021)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "CNNs have small receptive fields and fail to capture long-range global context across B-scans."),
        ("Dataset(s)", "UCSD and private scans."),
        ("Model", "Vision Transformer (ViT) with self-attention."),
        ("Explainability used?", "Yes (attention roll-out maps showing token importance)."),
        ("External validation?", "No."),
        ("Strengths", "Captures global structures; provides intrinsic attention explanations."),
        ("Weaknesses", "High computational footprint; prone to overfitting without massive pre-training."),
        ("What can our paper learn?", "ViT is a strong baseline, which is why we included it in our multi-baseline benchmarking table.")
    ])

    # 12. ReLaX (2022)
    doc.add_heading("12. ReLaX: Retinal Layer Attribution (2022)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Pixel-level attributions are clinically uninterpretable for doctors who evaluate in terms of layer boundaries."),
        ("Dataset(s)", "Duke OCT cohort and public sets."),
        ("Model", "ReLaX framework (classification + layer segmentation)."),
        ("Explainability used?", "Yes, computes attributions directly for anatomically defined retinal layers."),
        ("External validation?", "Yes."),
        ("Strengths", "Connects saliency maps directly to clinical terminology (e.g., ONL)."),
        ("Weaknesses", "Highly dependent on the segmentation sub-network; segmentation failure degrades explainability."),
        ("What can our paper learn?", "Clinicians trust layer localization. We use spatial attention to constrain representations without segmentations.")
    ])

    # 13. Score-CAM (2020)
    doc.add_heading("13. Score-CAM (2020)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Gradient-based attribution methods (like Grad-CAM) suffer from gradient noise and saturation."),
        ("Dataset(s)", "ImageNet, PASCAL VOC."),
        ("Model", "Standard CNNs (VGG, ResNet)."),
        ("Explainability used?", "Yes, proposed Score-CAM (weights activations by score changes without gradients)."),
        ("External validation?", "Yes."),
        ("Strengths", "Bypasses gradient noise; yields cleaner, less noisy attributions."),
        ("Weaknesses", "Computationally expensive (requires C forward passes, where C is the channel count)."),
        ("What can our paper learn?", "Serves as a vital benchmark for explaining gradient saturation issues in medical networks.")
    ])

    # 14. Integrated Gradients (2017)
    doc.add_heading("14. Integrated Gradients (2017)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Many saliency maps violate fundamental explanation axioms (Completeness and Invariance)."),
        ("Dataset(s)", "MNIST, ImageNet."),
        ("Model", "Inception, ResNet."),
        ("Explainability used?", "Yes, integrates gradients along a path from a reference image to the input."),
        ("External validation?", "Yes."),
        ("Strengths", "Mathematically proven to satisfy core explanation axioms."),
        ("Weaknesses", "Requires choosing an arbitrary reference baseline image; slow due to computing many path steps."),
        ("What can our paper learn?", "The baseline reference concept is used directly in our Insertion Test in evaluation.py.")
    ])

    # 15. Evaluation of Explainable Methods (2022)
    doc.add_heading("15. Evaluation of Explainable Methods (2022)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Lack of consensus on which explainability method is most reliable for clinical ophthalmic use."),
        ("Dataset(s)", "Multi-center clinical fundus and OCT datasets."),
        ("Model", "ResNet-50, VGG-16."),
        ("Explainability used?", "Yes (benchmarks Grad-CAM, Integrated Gradients, SHAP, LRP)."),
        ("External validation?", "Yes."),
        ("Strengths", "Compares qualitative physician ratings against quantitative perturbation tests."),
        ("Weaknesses", "Reveals that explanation methods frequently disagree, highlighting the vulnerability of visual attributions."),
        ("What can our paper learn?", "Confirms visual explanations are unstable, setting up our core research problem: auditing explainability decay under shifts.")
    ])

    # 16. RETFound (Zhou et al., Nature 2023)
    doc.add_heading("16. RETFound: Retinal Foundation Model (Nature 2023)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Training medical deep networks requires massive labeled datasets, which are scarce."),
        ("Dataset(s)", "Pre-trained on 1.6 million unlabelled fundus and OCT images."),
        ("Model", "ViT-based Masked Autoencoder (MAE)."),
        ("Explainability used?", "Yes (self-attention weights)."),
        ("External validation?", "Yes, across multiple unseen clinical external databases."),
        ("Strengths", "State-of-the-art generalization; requires minimal labels to fine-tune."),
        ("Weaknesses", "Massive computational cost to pre-train; requires fine-tuning on target datasets."),
        ("What can our paper learn?", "Serves as our primary benchmark for the future of retinal models, which we discuss in the future extensions section.")
    ])

    # 17. External Validation in Ophthalmology (2023)
    doc.add_heading("17. External Validation in Ophthalmology (2023)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Many papers claim clinical readiness but degrade significantly when tested on external clinics."),
        ("Dataset(s)", "Meta-analysis of ophthalmic AI publications."),
        ("Model", "Various."),
        ("Explainability used?", "No."),
        ("External validation?", "Yes (synthesizes findings across 50+ papers)."),
        ("Strengths", "Quantifies the typical accuracy drop (10-15%) caused by scanner domain shifts."),
        ("Weaknesses", "Evaluates only accuracy decay, ignoring explanation degradation."),
        ("What can our paper learn?", "Establishes the clinical necessity of our external validation on the OCTID dataset.")
    ])

    # 18. From Public to Clinical Data (MedViT, 2024)
    doc.add_heading("18. From Public to Clinical Data (2024)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Medical Vision Transformers (MedViT) degrade when moving from public datasets to clinical data."),
        ("Dataset(s)", "Public OCT datasets + real clinical cohort."),
        ("Model", "MedViT."),
        ("Explainability used?", "Yes."),
        ("External validation?", "Yes, focused on cross-dataset validation."),
        ("Strengths", "Highlights domain generalization limits of transformers in real clinics."),
        ("Weaknesses", "Does not offer a lightweight, zero-retraining adaptation method like our proposed normalization."),
        ("What can our paper learn?", "Strengthens our Discussion section on how domain shift affects transformers compared to attention-gated CNNs.")
    ])

    # 19. Calibration of Modern Neural Networks (Guo et al., 2017)
    doc.add_heading("19. Calibration of Modern Neural Networks (2017)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Modern deep neural networks are highly accurate but poorly calibrated."),
        ("Dataset(s)", "ImageNet, CIFAR."),
        ("Model", "ResNet, DenseNet."),
        ("Explainability used?", "No."),
        ("External validation?", "Yes."),
        ("Strengths", "Exposed the calibration problem; proposed Temperature Scaling as a simple fix."),
        ("Weaknesses", "Standard calibration methods do not account for calibration degradation under domain shifts."),
        ("What can our paper learn?", "Adopted the Expected Calibration Error (ECE) metric and binning logic implemented in evaluation.py.")
    ])

    # 20. TRIPOD-AI & CONSORT-AI Guidelines
    doc.add_heading("20. TRIPOD-AI & CONSORT-AI Guidelines", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Lack of standardized reporting guidelines for clinical AI, leading to poor reproducibility."),
        ("Dataset(s)", "Consensus guidelines from global clinical and technical panels."),
        ("Model", "N/A."),
        ("Explainability used?", "Yes (recommends declaring explanation protocols)."),
        ("External validation?", "Yes (declares external validation as a strict clinical requirement)."),
        ("Strengths", "Global reporting standard for clinical trials of AI models."),
        ("Weaknesses", "Policy and reporting guideline, not a technical implementation."),
        ("What can our paper learn?", "We structure our paper to meet TRIPOD-AI guidelines by declaring dataset splits and external validation.")
    ])

    # SECTION C: Newly Added Missing Literature Papers
    doc.add_page_break()
    doc.add_heading("Section 3: Newly Added Validation Papers", level=1)

    # 21. ROAD: Remove and Debias (2022)
    doc.add_heading("21. ROAD: Remove and Debias (Rong et al., 2022)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Standard pixel deletion tests introduce out-of-distribution (OOD) artifacts (sharp edges) that distort faithfulness scores."),
        ("Dataset(s)", "ImageNet, MNIST."),
        ("Model", "ResNet-50, VGG-16."),
        ("Explainability used?", "Yes, proposed the ROAD evaluation metric."),
        ("External validation?", "N/A (methodological paper)."),
        ("Strengths", "Replaces masked pixels using Noisy Linear Imputation, preserving distribution consistency."),
        ("Weaknesses", "The imputation model is simple (local linear) and might fail to preserve complex, high-res structures."),
        ("What can our paper learn?", "Validates our choice of using local contrast preprocessing to align data distributions during perturbation tests.")
    ])

    # 22. Quantus (Hedström et al., 2023)
    doc.add_heading("22. Quantus Evaluation Framework (2023)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Saliency evaluation is fragmented and poorly standardized, leading to reproducibility failures."),
        ("Dataset(s)", "MNIST, ImageNet."),
        ("Model", "Various CNNs."),
        ("Explainability used?", "Yes, provides a standardized Python framework containing 30+ XAI metrics."),
        ("External validation?", "N/A."),
        ("Strengths", "Standardizes XAI evaluation metrics into a unified, open-source library."),
        ("Weaknesses", "Primarily designed for 2D classification, requiring customization for 3D or segmented tasks."),
        ("What can our paper learn?", "Provides the scientific justification for evaluating multiple XAI dimensions in medical diagnostics.")
    ])

    # 23. Sanity Checks for Saliency Maps (Adebayo et al., 2018)
    doc.add_heading("23. Sanity Checks for Saliency Maps (2018)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Saliency maps are assumed to explain model reasoning, but they may simply act as edge-detectors, independent of weights."),
        ("Dataset(s)", "ImageNet, MNIST."),
        ("Model", "InceptionV3, ResNet-50."),
        ("Explainability used?", "Yes, evaluates Grad-CAM, Guided Backprop, and Integrated Gradients."),
        ("External validation?", "N/A."),
        ("Strengths", "Proposes Cascading Weight Randomization to test if heatmaps change when model weights are randomized."),
        ("Weaknesses", "Highlights that many popular XAI methods fail basic sanity checks."),
        ("What can our paper learn?", "Proves Grad-CAM can be misleading. Explains why we must use LayerCAM and mathematically verify it via AOPC.")
    ])

    # 24. ROAR: Remove and Retrain (Hooker et al., 2019)
    doc.add_heading("24. ROAR: Remove and Retrain (2019)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Deleting pixels in test images introduces covariate shift. Saliency drops could be due to shifted inputs."),
        ("Dataset(s)", "ImageNet, CIFAR-10."),
        ("Model", "ResNet-50."),
        ("Explainability used?", "Yes, proposed the ROAR evaluation framework."),
        ("External validation?", "N/A."),
        ("Strengths", "Retraining the model from scratch on modified images guarantees the model sees valid data distributions."),
        ("Weaknesses", "Extremely expensive computationally (requires retraining models hundreds of times for each step)."),
        ("What can our paper learn?", "ROAR is the gold standard, but AOPC serves as a computationally feasible alternative that we debias via preprocessing.")
    ])

    # 25. DomainBed (Gulrajani & Lopez-Paz, 2021)
    doc.add_heading("25. DomainBed (2021)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Benchmark algorithms for domain generalization are evaluated inconsistently with cherry-picked hyperparameter tuning."),
        ("Dataset(s)", "PACS, VLCS, OfficeHome."),
        ("Model", "ResNet-50 backbone."),
        ("Explainability used?", "No."),
        ("External validation?", "Yes (focuses entirely on out-of-distribution domain generalization)."),
        ("Strengths", "Standardizes domain generalization evaluations under strict model selection rules."),
        ("Weaknesses", "Shows that simple Empirical Risk Minimization (ERM) often matches or outperforms complex domain adaptation models."),
        ("What can our paper learn?", "Supports our design choice of using deterministic, input-level standardization (CLAHE + Min-Max) over adaptation losses.")
    ])

    # 26. Reliability Diagrams (Niculescu-Mizil & Caruana, 2005)
    doc.add_heading("26. Reliability Diagrams (2005)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Need for graphical representations of model calibration to identify underconfidence and overconfidence."),
        ("Dataset(s)", "Synthetic, text, and image datasets."),
        ("Model", "SVMs, Naive Bayes, Neural Networks."),
        ("Explainability used?", "No."),
        ("External validation?", "Yes."),
        ("Strengths", "Introduced visual accuracy-confidence bin plots ($y=x$ mapping) to diagnose model miscalibration."),
        ("Weaknesses", "Does not provide a single scalar score, requiring complementation by quantitative ECE."),
        ("What can our paper learn?", "We combine ECE calculation with reliability diagrams to track calibration degradation under domain shift.")
    ])

    # 27. DECIDE-AI Guidelines (Vasey et al., 2022)
    doc.add_heading("27. DECIDE-AI Guidelines (Vasey et al., 2022)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Translational gap in clinical AI. Models with high laboratory accuracy fail when deployed in clinical trials."),
        ("Dataset(s)", "N/A (consensus statement)."),
        ("Model", "N/A."),
        ("Explainability used?", "Yes (recommends declaring explanation protocols)."),
        ("External validation?", "Yes (declares external validation as a strict clinical requirement)."),
        ("Strengths", "First international reporting guideline for clinical AI systems."),
        ("Weaknesses", "Policy guideline, not a technical implementation."),
        ("What can our paper learn?", "Provides the clinical justification for our study's scope: evaluating explanation faithfulness computationally.")
    ])

    # 28. Scanner Intensity Shift in OCT (Gholami et al., 2020)
    doc.add_heading("28. Scanner Intensity Shift in OCT (2020)", level=2)
    add_styled_table(doc, ["Question", "Your Notes"], [
        ("Problem addressed", "Performance decay of retinal classifiers when tested on scans from different scanner models (Spectralis vs. Cirrus)."),
        ("Dataset(s)", "Multi-vendor clinical OCT cohorts."),
        ("Model", "DenseNet, ResNet."),
        ("Explainability used?", "Yes, cross-scanner validation."),
        ("Strengths", "Proves that differences in contrast and speckle noise cause feature representation shifts."),
        ("Weaknesses", "Relies on retraining models on mixed-scanner data, which is not possible in single-scanner clinics."),
        ("What can our paper learn?", "Directly motivates Project 2: using input-level normalizations (CLAHE + Min-Max) to align dynamic ranges.")
    ])

    # Save Document
    filename = "researchrelated details.docx"
    doc.save(filename)
    print(f"Successfully generated MS Word file: {os.path.abspath(filename)}")

if __name__ == "__main__":
    main()
